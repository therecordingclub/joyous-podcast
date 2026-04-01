#!/usr/bin/env python3
"""
Generate the Joyous Podcast Growth & Distribution Strategy document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color palette ──────────────────────────────────────────────────────
DEEP_TEAL   = RGBColor(0x0C, 0x4B, 0x5E)  # Primary headings
WARM_GOLD   = RGBColor(0xC8, 0x96, 0x2E)  # Accent
SOFT_CREAM  = RGBColor(0xFA, 0xF5, 0xEB)  # Background accents
DARK_GRAY   = RGBColor(0x2D, 0x2D, 0x2D)  # Body text
MED_GRAY    = RGBColor(0x5A, 0x5A, 0x5A)  # Secondary text
LIGHT_TEAL  = RGBColor(0xE8, 0xF4, 0xF7)  # Table header bg
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Custom styles ──────────────────────────────────────────────────────
style = doc.styles

# -- Body text
normal = style['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK_GRAY
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# -- Heading 1 (chapter titles)
h1 = style['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(26)
h1.font.bold = True
h1.font.color.rgb = DEEP_TEAL
h1.paragraph_format.space_before = Pt(0)
h1.paragraph_format.space_after = Pt(16)
h1.paragraph_format.keep_with_next = True
# Add bottom border to H1
h1_pPr = h1.element.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                 f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="C8962E"/>'
                 f'</w:pBdr>')
h1_pPr.append(pBdr)

# -- Heading 2 (section titles)
h2 = style['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = DEEP_TEAL
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)

# -- Heading 3
h3 = style['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = WARM_GOLD
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)

# -- List Bullet
lb = style['List Bullet']
lb.font.name = 'Calibri'
lb.font.size = Pt(11)
lb.font.color.rgb = DARK_GRAY
lb.paragraph_format.space_after = Pt(3)

# ── Helper functions ───────────────────────────────────────────────────

def add_page_break():
    doc.add_page_break()

def add_bold_para(text, size=11, color=DARK_GRAY, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_para(text, bold_prefix=None, size=11):
    p = doc.add_paragraph()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(size)
        run_b.font.color.rgb = DARK_GRAY
        run_b.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level == 1:
        p.paragraph_format.left_indent = Cm(2.0)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = DARK_GRAY
        run_b.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    return p

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = DARK_GRAY
        run_b.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '0C4B5E')
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Calibri'
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F5F5F5')
    doc.add_paragraph()  # spacer
    return table

def add_callout_box(text, title=None):
    """Add a highlighted callout box using a bordered paragraph."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    # Add border and shading
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="C8962E"/>'
        f'</w:pBdr>'
    )
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FAF5EB" w:val="clear"/>')
    pPr.append(borders)
    pPr.append(shading)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    if title:
        run_t = p.add_run(title + "  ")
        run_t.bold = True
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = WARM_GOLD
        run_t.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════

# Top spacer
for _ in range(6):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('JOYOUS')
run.font.size = Pt(52)
run.font.bold = True
run.font.color.rgb = DEEP_TEAL
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PODCAST')
run.font.size = Pt(36)
run.font.bold = False
run.font.color.rgb = WARM_GOLD
run.font.name = 'Calibri'
run.font.italic = True

# Horizontal rule
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.color.rgb = WARM_GOLD
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Growth & Distribution Strategy')
run.font.size = Pt(22)
run.font.color.rgb = DEEP_TEAL
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Health & Wealth at the Intersection of East and West')
run.font.size = Pt(14)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'
run.font.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for Jodi Yang')
run.font.size = Pt(14)
run.font.color.rgb = DARK_GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('March 2026')
run.font.size = Pt(12)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONFIDENTIAL')
run.font.size = Pt(10)
run.font.color.rgb = WARM_GOLD
run.font.name = 'Calibri'
run.bold = True


# ═══════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('Executive Summary', '3'),
    ('1. YouTube Strategy', '4'),
    ('2. Podcast Distribution Strategy', '8'),
    ('3. Bilingual Release Strategy', '12'),
    ('4. Social Media Strategy', '16'),
    ('5. Prime Lab Show Concept', '20'),
    ('6. Placement Agency Business Strategy', '23'),
    ('7. Monetization & Revenue Streams', '27'),
    ('8. 90-Day Launch Plan', '30'),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.color.rgb = DEEP_TEAL
    run.font.name = 'Calibri'
    if title == 'Executive Summary':
        run.bold = True
    # Add tab and page number
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(13)
    run2.font.color.rgb = MED_GRAY
    run2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(8)
    # Add dotted tab leader
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)


# ═══════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('Executive Summary', level=1)

add_para(
    'The Joyous Podcast sits at a rare and powerful intersection: the convergence of health, '
    'wealth, Eastern and Western perspectives, and elite professional networks. With six edited '
    'episodes already produced, host Jodi Yang brings an unparalleled combination of '
    'credentials\u2014Wharton and Harvard education, senior roles at global investment firms from '
    'Citi to Balderton Capital, and genuine bilingual fluency in English and Mandarin Chinese.'
)

add_para(
    'This strategy document provides a comprehensive playbook for scaling Joyous from a '
    'promising early-stage podcast into a multi-platform media brand that simultaneously '
    'supports Jodi\'s professional ventures: a venture capital placement agency targeting '
    'global LPs, the Prime Lab F1 racing show concept, and her role as venture partner at '
    'Prime Movers Lab.'
)

doc.add_heading('Core Strategic Thesis', level=2)

add_para(
    'Joyous is not just a podcast\u2014it is the content engine for a vertically integrated media '
    'and financial services platform. Every piece of content serves dual purposes:'
)

add_bullet('Builds audience trust and brand equity in the health-and-wealth space', bold_prefix='Audience Growth: ')
add_bullet('Creates a pipeline of LP relationships for the placement agency', bold_prefix='Business Development: ')
add_bullet('Positions Jodi as the go-to connector between East and West in venture capital', bold_prefix='Network Leverage: ')
add_bullet('Opens sponsorship, event, and consulting revenue streams', bold_prefix='Monetization: ')

doc.add_heading('Unique Differentiators', level=2)

add_bullet('Only bilingual (English/Mandarin) health-and-wealth podcast targeting global audiences')
add_bullet('Host has direct investment experience across 3 continents and 12+ cities')
add_bullet('Natural bridge between Western venture capital and Asian LP capital')
add_bullet('Classical music background adds cultural depth and differentiation')
add_bullet('Prime Movers Lab affiliation provides deep-tech credibility and guest pipeline')

doc.add_heading('Key Recommendations at a Glance', level=2)

add_styled_table(
    ['Area', 'Priority Action', 'Timeline'],
    [
        ['YouTube', 'Launch bilingual channel with Shorts-first strategy', 'Weeks 1-4'],
        ['Podcast', 'Distribute to 10+ platforms including Chinese apps', 'Weeks 1-2'],
        ['Bilingual', 'Release dual-language episodes; launch Xiaohongshu', 'Weeks 3-6'],
        ['Social', 'LinkedIn thought leadership + TikTok/Reels clips', 'Ongoing'],
        ['Prime Lab', 'Produce pilot episode; secure 2 sponsors', 'Months 2-3'],
        ['Placement', 'File SEC Form D; onboard 3 GP clients', 'Months 1-3'],
        ['Revenue', 'Launch Patreon/membership tier', 'Month 2'],
    ]
)


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 1 — YOUTUBE STRATEGY
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('1. YouTube Strategy', level=1)

add_para(
    'YouTube is the single most important distribution channel for Joyous. It is the world\'s '
    'second-largest search engine, offers the best long-form discovery algorithm, and provides '
    'direct monetization. For a bilingual podcast bridging health and wealth across cultures, '
    'YouTube\'s global reach is unmatched.'
)

# -- Channel Setup & Branding
doc.add_heading('Channel Setup & Branding', level=2)

add_para('The channel name should reflect the bilingual nature and warmth of the brand:')

doc.add_heading('Bilingual Channel Name Ideas', level=3)

add_styled_table(
    ['English Name', 'Chinese Name', 'Notes'],
    [
        ['Joyous with Jodi Yang', 'Joyous \u6b22\u4e50\u64ad\u5ba2', 'Simple, personal, bilingual'],
        ['Joyous \u00b7 Health & Wealth', 'Joyous \u00b7 \u5065\u5eb7\u4e0e\u8d22\u5bcc', 'Descriptive, SEO-friendly'],
        ['The Joyous Podcast \u559c\u60a6\u64ad\u5ba2', '\u559c\u60a6\u64ad\u5ba2', 'Elegant Chinese translation'],
        ['Joyous \u2014 East Meets Wealth', 'Joyous \u2014 \u4e1c\u897f\u65b9\u8d22\u5bcc\u5bf9\u8bdd', 'Emphasizes cross-cultural angle'],
    ]
)

add_callout_box(
    'Use "Joyous with Jodi Yang" as the primary channel name and "\u559c\u60a6\u64ad\u5ba2 | Joyous Podcast" '
    'as the Chinese-facing handle. This allows the algorithm to surface the channel to both '
    'English and Chinese-speaking audiences.',
    title='RECOMMENDATION:'
)

doc.add_heading('Channel Art Requirements', level=3)

add_bullet('Banner: 2560x1440px, showing Jodi with bilingual text and warm gold/teal palette')
add_bullet('Profile photo: Professional headshot with warm tones, consistent across all platforms')
add_bullet('Channel trailer: 60-90 second sizzle reel mixing English and Mandarin clips')
add_bullet('Channel description: Full bilingual description with keywords in both languages')
add_bullet('Links: Website, all social profiles, Xiaohongshu, WeChat article link')

# -- Content Formats
doc.add_heading('Content Format Recommendations', level=2)

add_styled_table(
    ['Format', 'Length', 'Frequency', 'Purpose'],
    [
        ['Full Episode', '45-75 min', '1/week', 'Deep-dive interviews; core content'],
        ['Episode Highlights', '8-15 min', '2/week', 'Best segments; drives full episode views'],
        ['YouTube Shorts', '30-59 sec', '4-5/week', 'Viral hooks; algorithm growth'],
        ['Chinese Subtitled', '45-75 min', '2/month', 'Full episodes for Chinese audiences'],
        ['Solo Insights', '5-10 min', '2/month', 'Jodi\'s personal takes; builds authority'],
        ['Behind the Scenes', '2-5 min', '1/month', 'Humanizes brand; travel/lifestyle'],
    ]
)

# -- Thumbnails
doc.add_heading('Thumbnail Design Guidelines', level=2)

add_para('Thumbnails are the single biggest driver of click-through rate. Follow these rules:')

add_numbered('Use a close-up face showing genuine emotion (surprise, laughter, intensity)', bold_prefix='Faces sell: ')
add_numbered('Maximum 4-5 words in large, bold sans-serif font', bold_prefix='Minimal text: ')
add_numbered('Deep teal background + warm gold text = Joyous brand identity', bold_prefix='Brand colors: ')
add_numbered('Maintain 70%+ contrast ratio between text and background', bold_prefix='High contrast: ')
add_numbered('For bilingual episodes, put Chinese text on the left (read first in Chinese)', bold_prefix='Bilingual text: ')
add_numbered('Test 3 thumbnail variants using YouTube A/B testing feature', bold_prefix='A/B test: ')

# -- SEO
doc.add_heading('SEO Optimization: English & Chinese', level=2)

add_para(
    'Every video should be optimized for discovery in both languages. YouTube\'s algorithm '
    'indexes titles, descriptions, tags, and closed captions.'
)

doc.add_heading('English Keywords to Target', level=3)
add_bullet('health and wealth podcast, health wealth intersection, longevity and finance')
add_bullet('venture capital podcast, LP investing, fundraising strategy')
add_bullet('Asian American entrepreneur, bilingual business podcast')
add_bullet('wellness investing, biohacking finance, health tech investing')

doc.add_heading('Chinese Keywords to Target', level=3)
add_bullet('\u5065\u5eb7\u4e0e\u8d22\u5bcc (health and wealth), \u98ce\u9669\u6295\u8d44\u64ad\u5ba2 (venture capital podcast)')
add_bullet('\u534e\u4eba\u521b\u4e1a (Chinese entrepreneurship), \u4e1c\u897f\u65b9\u6295\u8d44 (East-West investing)')
add_bullet('\u5065\u5eb7\u79d1\u6280 (health tech), \u751f\u7269\u79d1\u6280\u6295\u8d44 (biotech investing)')
add_bullet('\u7f8e\u56fd\u98ce\u6295 (US venture capital), LP\u6295\u8d44\u7b56\u7565 (LP investment strategy)')

doc.add_heading('Description Template', level=3)
add_para(
    'Every video description should follow this structure: (1) compelling 2-sentence hook with '
    'primary keyword, (2) 3-4 sentence summary, (3) timestamps for all segments, (4) guest bio '
    'with links, (5) Chinese language summary paragraph, (6) social links and calls to action, '
    '(7) 10-15 hashtags in both languages.'
)

# -- Upload Schedule
doc.add_heading('Upload Schedule', level=2)

add_styled_table(
    ['Day', 'Content Type', 'Time (US)', 'Time (China)'],
    [
        ['Monday', 'YouTube Short', '8:00 AM ET', '8:00 PM CST'],
        ['Tuesday', 'Full Episode', '9:00 AM ET', '9:00 PM CST'],
        ['Wednesday', 'YouTube Short', '8:00 AM ET', '8:00 PM CST'],
        ['Thursday', 'Highlights Clip', '9:00 AM ET', '9:00 PM CST'],
        ['Friday', 'YouTube Short', '8:00 AM ET', '8:00 PM CST'],
        ['Saturday', 'YouTube Short (Chinese)', '10:00 AM ET', '10:00 PM CST'],
    ]
)

add_callout_box(
    'Post Shorts at 8 AM ET / 8 PM China time to catch both US morning scroll and '
    'Chinese evening scroll. Full episodes on Tuesday avoids the Monday content glut '
    'and Friday dropoff.',
    title='TIMING:'
)

# -- Shorts Strategy
doc.add_heading('YouTube Shorts Strategy for Viral Growth', level=2)

add_para(
    'Shorts are the #1 growth lever on YouTube in 2026. They feed the algorithm, attract '
    'subscribers, and cost almost nothing to produce from existing long-form content.'
)

add_bullet('Extract the single most surprising or emotional 45-second clip from each episode', bold_prefix='Hook-first clips: ')
add_bullet('Caption every Short in both English and Chinese (burned-in captions)', bold_prefix='Dual-language captions: ')
add_bullet('End every Short with a spoken call-to-action: "Full episode in bio"', bold_prefix='CTA: ')
add_bullet('Create "health myth vs. reality" and "money myth vs. reality" series', bold_prefix='Series format: ')
add_bullet('Use trending sounds sparingly but strategically when they fit the brand', bold_prefix='Trend-jack: ')
add_bullet('Repost top-performing Shorts to TikTok and Instagram Reels simultaneously', bold_prefix='Cross-post: ')

# -- Monetization
doc.add_heading('Monetization Timeline', level=2)

add_styled_table(
    ['Milestone', 'Requirement', 'Est. Timeline', 'Revenue Potential'],
    [
        ['YouTube Partner Program', '1,000 subs + 4,000 watch hrs', 'Months 3-6', '$200-500/mo'],
        ['Shorts Monetization', '1,000 subs + 10M Short views', 'Months 4-8', '$100-300/mo'],
        ['Sponsored Videos', '5,000+ subs, niche authority', 'Months 4-6', '$500-2,000/video'],
        ['Channel Memberships', '1,000+ subs', 'Month 6', '$300-800/mo'],
        ['Super Chats (Live)', 'Partner Program active', 'Month 6+', '$50-200/stream'],
    ]
)

# -- Collaboration
doc.add_heading('Collaboration Strategy', level=2)

add_para('Strategic collaborations accelerate growth 3-5x compared to organic posting alone.')

doc.add_heading('Target Collaborators', level=3)
add_bullet('Health/longevity: Andrew Huberman (guest pitch), Peter Attia, Rhonda Patrick')
add_bullet('Wealth/VC: All-In Podcast (guest pitch), 20VC with Harry Stebbings, The Tim Ferriss Show')
add_bullet('Bilingual/Asian American: The Wong Fu crew, Simu Liu, Chinese-American creators')
add_bullet('Finance education: Graham Stephan, Andrei Jikh, Humphrey Yang')
add_bullet('Chinese podcasters: \u65e5\u8c08\u516c\u56ed (Ritangongyuan), \u5c0f\u5b87\u5b99 (Xiaoyuzhou) top hosts')

doc.add_heading('Collaboration Formats', level=3)
add_bullet('Guest swaps: Appear on their show, invite them to Joyous')
add_bullet('Joint Shorts: Co-create 3-5 Shorts during a single recording session')
add_bullet('Live events: Co-host a YouTube Live on a trending health/wealth topic')
add_bullet('Playlist features: Create curated playlists featuring complementary creators')

# -- KPIs
doc.add_heading('Analytics KPIs to Track', level=2)

add_styled_table(
    ['Metric', 'Target (Month 3)', 'Target (Month 6)', 'Target (Year 1)'],
    [
        ['Subscribers', '1,000', '5,000', '25,000'],
        ['Avg. View Duration', '40%+', '45%+', '50%+'],
        ['Click-Through Rate', '4%+', '6%+', '8%+'],
        ['Shorts Views/Month', '50,000', '500,000', '2,000,000'],
        ['Watch Hours/Month', '500', '2,000', '10,000'],
        ['Revenue/Month', '$0', '$500', '$3,000+'],
    ]
)


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 2 — PODCAST DISTRIBUTION
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('2. Podcast Distribution Strategy', level=1)

add_para(
    'With six edited episodes ready, the immediate priority is comprehensive distribution. '
    'The podcast must be available everywhere listeners already consume content\u2014across Western '
    'and Chinese platforms.'
)

# -- Platform Distribution
doc.add_heading('Platform Distribution Matrix', level=2)

add_styled_table(
    ['Platform', 'Priority', 'Audience', 'Notes'],
    [
        ['Apple Podcasts', 'Critical', 'US/Global English', '#1 for iPhone users; strong in US'],
        ['Spotify', 'Critical', 'US/Global', '#1 overall; video podcast support growing'],
        ['YouTube Music', 'Critical', 'Global', 'Auto-syncs from YouTube; growing fast'],
        ['Amazon Music / Audible', 'High', 'US', 'Growing market share; Alexa integration'],
        ['iHeartRadio', 'High', 'US', 'Large US audience; radio crossover'],
        ['Pocket Casts', 'Medium', 'Tech-savvy', 'Popular with tech/business audience'],
        ['Overcast', 'Medium', 'iOS power users', 'Influential early adopters'],
        ['Castro', 'Medium', 'iOS', 'Niche but loyal audience'],
        ['Podcast Addict', 'Medium', 'Android', '#1 Android podcast app'],
        ['Castbox', 'Medium', 'Global/Asian', 'Strong in Asian markets'],
        ['Podbean', 'Low', 'Global', 'Hosting + distribution'],
        ['TuneIn', 'Low', 'Smart speakers', 'Alexa/Google Home discovery'],
    ]
)

# -- RSS Feed
doc.add_heading('RSS Feed Setup', level=2)

add_para(
    'A single RSS feed powers distribution to all Western platforms. Choose a hosting provider '
    'that supports both audio and video podcasting:'
)

doc.add_heading('Recommended Hosting Platforms', level=3)

add_styled_table(
    ['Host', 'Price', 'Best For', 'Key Feature'],
    [
        ['Riverside.fm', '$15-24/mo', 'Recording + hosting', 'High-quality remote recording'],
        ['Buzzsprout', '$12-24/mo', 'Simplicity', 'Auto-distribution to all platforms'],
        ['Transistor', '$19-49/mo', 'Analytics', 'Multiple shows under one account'],
        ['Libsyn', '$5-20/mo', 'Budget', 'Industry standard; reliable'],
        ['Anchor (Spotify)', 'Free', 'Beginners', 'Free but Spotify-owned (lock-in risk)'],
    ]
)

add_callout_box(
    'Use Buzzsprout or Transistor for the main English feed. Both auto-submit to Apple Podcasts, '
    'Spotify, Amazon, and 10+ other directories with one click. Transistor is ideal if you plan '
    'to run multiple shows (e.g., Joyous + Prime Lab as separate feeds).',
    title='RECOMMENDATION:'
)

# -- Chinese Platforms
doc.add_heading('Chinese Podcast Platforms', level=2)

add_para(
    'China\'s podcast market reached $3.2 billion in 2025 and is growing 25% year-over-year. '
    'Jodi\'s Mandarin fluency is a massive competitive advantage\u2014almost no Western-based '
    'podcasters distribute natively on Chinese platforms.'
)

add_styled_table(
    ['Platform', 'Chinese Name', 'Monthly Users', 'Strategy'],
    [
        ['Ximalaya', '\u559c\u9a6c\u62c9\u96c5', '300M+', '#1 audio platform; must-have; supports paid content'],
        ['Lizhi', '\u8354\u679d', '50M+', 'Live audio + podcast; younger demographic'],
        ['QingTing FM', '\u873b\u8713FM', '100M+', 'Strong in commuter/car audio'],
        ['Apple Podcasts China', 'Apple \u64ad\u5ba2', '30M+', 'Separate submission; same RSS feed works'],
        ['Xiaoyuzhou', '\u5c0f\u5b87\u5b99', '10M+', 'Premium indie podcast app; tastemaker audience'],
        ['NetEase Cloud Music', '\u7f51\u6613\u4e91\u97f3\u4e50', '200M+', 'Music + podcast integration'],
        ['Bilibili', '\u54d4\u54e9\u54d4\u54e9', '350M+', 'Video podcast; Gen Z audience; high engagement'],
    ]
)

add_callout_box(
    'Start with Ximalaya (\u559c\u9a6c\u62c9\u96c5) and Xiaoyuzhou (\u5c0f\u5b87\u5b99) as the two priority Chinese platforms. '
    'Ximalaya gives mass reach; Xiaoyuzhou gives credibility with China\'s podcast-savvy early '
    'adopters. Both require separate account registration and content upload (RSS auto-import '
    'is limited for Chinese platforms).',
    title='CHINA PRIORITY:'
)

# -- Cross-Promotion
doc.add_heading('Cross-Promotion Strategies', level=2)

add_numbered('Include a "where to listen" page at joyouspodcast.com with links to every platform', bold_prefix='Smart Links: ')
add_numbered('End every episode with "subscribe on [platform] and leave a review"', bold_prefix='Verbal CTAs: ')
add_numbered('Create a WeChat mini-program or article linking to Chinese platform episodes', bold_prefix='WeChat Integration: ')
add_numbered('When guests appear, provide them a pre-written social post with listen links', bold_prefix='Guest Amplification: ')
add_numbered('Join podcast networks or cross-promotion platforms like Podchaser or Chartable', bold_prefix='Network Effects: ')

# -- Reviews
doc.add_heading('Review & Rating Growth Tactics', level=2)

add_para(
    'Apple Podcasts reviews and ratings are critical for chart placement. Target 50 reviews '
    'within the first 30 days of launch.'
)

add_bullet('Ask listeners directly at the end of each episode (specific CTA with instructions)')
add_bullet('Create a "review and screenshot" giveaway (entry = review screenshot DM\'d to Instagram)')
add_bullet('Email personal network (Wharton alumni, Harvard MBA classmates) asking for launch-week reviews')
add_bullet('Offer a "founding listener" shoutout for the first 100 reviewers')
add_bullet('Create a QR code linking directly to the Apple Podcasts review page')

# -- Guest Booking
doc.add_heading('Guest Booking Strategy', level=2)

doc.add_heading('Ideal Guest Profiles', level=3)
add_bullet('Health-tech founders and executives (especially Prime Movers Lab portfolio)')
add_bullet('Longevity researchers and biohacking practitioners')
add_bullet('Wealth managers and family office principals (potential LP leads)')
add_bullet('Cross-cultural business leaders (US-China, US-Asia bridge figures)')
add_bullet('Classical musicians who crossed into business (unique angle)')
add_bullet('Angel investors and emerging fund managers')

doc.add_heading('Guest Outreach Template', level=3)

add_callout_box(
    'Subject: Guest Invitation \u2014 Joyous Podcast (Health x Wealth)\n\n'
    'Hi [Name],\n\n'
    'I\'m Jodi Yang, host of Joyous\u2014a podcast exploring the intersection of health and wealth. '
    'I\'m a venture partner at Prime Movers Lab with a background spanning Citi, KKR-acquired '
    'Novo Holdings, B Capital, and Balderton Capital.\n\n'
    'I\'d love to feature you on an upcoming episode to discuss [specific topic related to their '
    'expertise]. Our audience includes investors, founders, and health-conscious professionals '
    'across English and Chinese-speaking markets.\n\n'
    'The conversation would be 45-60 minutes, recorded remotely via Riverside. We handle all '
    'editing, promotion, and distribution across 15+ platforms including major Chinese audio apps.\n\n'
    'Would you be open to a 15-minute call to explore this?\n\n'
    'Warmly,\nJodi',
    title='EMAIL TEMPLATE:'
)

# -- Recording Quality
doc.add_heading('Recording Quality Standards', level=2)

add_bullet('Microphone: Shure SM7B or Rode PodMic (dynamic mics reduce room noise)', bold_prefix='Audio: ')
add_bullet('Interface: Focusrite Scarlett 2i2 or Rodecaster Pro II', bold_prefix='Interface: ')
add_bullet('Remote recording: Riverside.fm (uncompressed local recording, not Zoom)', bold_prefix='Remote: ')
add_bullet('Camera: Sony ZV-1 or iPhone 15 Pro with good lighting (two softboxes minimum)', bold_prefix='Video: ')
add_bullet('Room: Acoustic treatment (foam panels or blankets); avoid hard surfaces', bold_prefix='Environment: ')
add_bullet('Backup: Always record a local backup on phone or secondary device', bold_prefix='Redundancy: ')
add_bullet('Post-production: Descript for editing; Adobe Podcast for AI noise removal', bold_prefix='Editing: ')


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 3 — BILINGUAL RELEASE STRATEGY
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('3. Bilingual Release Strategy', level=1)

add_para(
    'Jodi\'s fluency in English and Mandarin Chinese is the podcast\'s most distinctive asset. '
    'No other health-and-wealth podcast credibly bridges both language markets. This chapter '
    'provides a framework for leveraging bilingualism without doubling the production workload.'
)

# -- Dual-language framework
doc.add_heading('Dual-Language Episode Framework', level=2)

add_styled_table(
    ['Episode Type', 'Language Mix', 'Frequency', 'Target Audience'],
    [
        ['Standard Interview', '100% English', '3/month', 'Global English speakers'],
        ['Bilingual Interview', '70% EN / 30% CN', '1/month', 'Bilingual diaspora; curious EN listeners'],
        ['Full Chinese Episode', '100% Mandarin', '1/month', 'Mainland China, Taiwan, Singapore'],
        ['Chinese Recap', '100% Mandarin, 10 min', '2/month', 'Summary of EN episodes for CN audience'],
        ['Code-Switch Shorts', 'Mixed EN/CN', '4/month', 'Social media; viral bilingual clips'],
    ]
)

add_callout_box(
    'The "Chinese Recap" format is the highest-ROI bilingual content: Jodi records a '
    '10-minute Mandarin summary of the best English episodes. Minimal production cost, '
    'maximum reach into Chinese-speaking markets. These also work perfectly as Bilibili '
    'and Xiaohongshu video content.',
    title='HIGH-ROI FORMAT:'
)

# -- When to do what
doc.add_heading('When to Use Each Format', level=2)

doc.add_heading('Full Chinese Episodes', level=3)
add_bullet('When the guest speaks Mandarin (Chinese founders, Asian family office heads)')
add_bullet('Topics with specific China relevance (Chinese health tech, TCM + modern medicine)')
add_bullet('During Chinese holidays (Lunar New Year special, Mid-Autumn Festival)')
add_bullet('When covering China-specific investment themes')

doc.add_heading('Bilingual Episodes', level=3)
add_bullet('When the guest is bilingual (natural code-switching feels authentic)')
add_bullet('When discussing cross-cultural topics (East vs. West health approaches)')
add_bullet('When a concept is better expressed in the other language')
add_bullet('For "language lesson" segments (teach a Chinese health/wealth concept)')

doc.add_heading('Subtitled English Episodes', level=3)
add_bullet('All full English episodes should get Chinese subtitles on YouTube')
add_bullet('Use AI-assisted translation (DeepL + human review) to keep costs down')
add_bullet('Burned-in Chinese subtitles for Bilibili uploads')
add_bullet('SRT files for YouTube (allows viewer choice)')

# -- Chinese Social Platforms
doc.add_heading('Chinese Social Media Platforms', level=2)

add_styled_table(
    ['Platform', 'Chinese Name', 'Format', 'Priority', 'Strategy'],
    [
        ['Xiaohongshu', '\u5c0f\u7ea2\u4e66', 'Photo + short video', 'Critical',
         'Health/wealth tips; lifestyle; cross-cultural insights'],
        ['WeChat Official Account', '\u5fae\u4fe1\u516c\u4f17\u53f7', 'Long articles', 'Critical',
         'Episode summaries; newsletter substitute for China'],
        ['Weibo', '\u5fae\u535a', 'Short text + media', 'High',
         'News/opinions; engagement with Chinese media'],
        ['Bilibili', '\u54d4\u54e9\u54d4\u54e9', 'Long video', 'High',
         'Full Chinese episodes + subtitled English episodes'],
        ['Douyin', '\u6296\u97f3', 'Short video', 'High',
         'Chinese version of TikTok; 60-sec health/wealth clips'],
    ]
)

doc.add_heading('Xiaohongshu (\u5c0f\u7ea2\u4e66) Deep Dive', level=3)

add_para(
    'Xiaohongshu (literally "Little Red Book") is the #1 lifestyle platform for educated, '
    'urban Chinese women aged 25-45\u2014exactly Joyous\'s target Chinese demographic. With 300M+ '
    'monthly active users, it is the Instagram + Pinterest of China.'
)

add_bullet('Post 3-4x per week: mix of carousel posts (health tips), short videos, and life updates')
add_bullet('Use trending Chinese hashtags: #\u5065\u5eb7\u751f\u6d3b #\u7406\u8d22\u77e5\u8bc6 #\u7f8e\u56fd\u751f\u6d3b #\u6295\u8d44\u7b14\u8bb0 #\u5973\u6027\u529b\u91cf')
add_bullet('Share behind-the-scenes of podcast recording (Chinese audiences love process content)')
add_bullet('Create "bilingual vocabulary" posts teaching business/health terms')
add_bullet('Cross-cultural comparison content performs extremely well (\u4e2d\u7f8e\u5bf9\u6bd4)')

doc.add_heading('WeChat Official Account Strategy', level=3)

add_para(
    'WeChat is essential for reaching Chinese professionals. An Official Account functions '
    'as a newsletter, blog, and distribution channel combined.'
)

add_bullet('Publish 2x per week: episode summaries + original bilingual content')
add_bullet('Build a WeChat group (\u5fae\u4fe1\u7fa4) for engaged listeners\u2014cap at 200 for quality')
add_bullet('Use WeChat mini-programs to link to podcast episodes on Ximalaya')
add_bullet('Share QR code at events and in email signatures')

# -- Translation Workflow
doc.add_heading('Translation & Localization Workflow', level=2)

add_numbered('Record episode in primary language', bold_prefix='Step 1: ')
add_numbered('Generate AI transcript using Descript or Whisper', bold_prefix='Step 2: ')
add_numbered('Run through DeepL for initial translation', bold_prefix='Step 3: ')
add_numbered('Native speaker review (hire a bilingual VA for $15-25/hr)', bold_prefix='Step 4: ')
add_numbered('Cultural adaptation: replace idioms, adjust references for target culture', bold_prefix='Step 5: ')
add_numbered('Create subtitles (SRT for YouTube, burned-in for Bilibili/Douyin)', bold_prefix='Step 6: ')
add_numbered('Write platform-specific social copy in target language', bold_prefix='Step 7: ')

add_callout_box(
    'Budget approximately $100-200 per episode for translation and localization. This includes '
    'AI-assisted translation ($0), native speaker review ($50-100), subtitle creation ($30-50), '
    'and social media copy adaptation ($20-50). As volume increases, negotiate a retainer with '
    'a bilingual content assistant.',
    title='COST ESTIMATE:'
)

# -- Cultural Adaptation
doc.add_heading('Cultural Adaptation Considerations', level=2)

add_bullet('Health topics: Chinese audiences are more receptive to TCM (Traditional Chinese Medicine) integration; Western audiences prefer evidence-based framing')
add_bullet('Wealth topics: Avoid discussing specific investment returns for China content (regulatory sensitivity); focus on strategy and philosophy')
add_bullet('Humor: Self-deprecating humor works in English; in Chinese, observational humor about cross-cultural differences resonates better')
add_bullet('Formality: Chinese podcast audiences expect slightly more formality in host introductions; use proper titles')
add_bullet('Guest introductions: In Chinese, lead with institutional affiliation before personal story')

# -- Target Demographics
doc.add_heading('Target Chinese-Speaking Demographics', level=2)

add_styled_table(
    ['Segment', 'Location', 'Language', 'Content Preference'],
    [
        ['Urban professionals', 'Mainland China', 'Mandarin', 'Wealth building, health tech, US insights'],
        ['Tech workers', 'Taiwan', 'Mandarin (traditional)', 'Startup culture, work-life balance'],
        ['Finance professionals', 'Singapore/HK', 'Mandarin + English', 'Cross-border investing, bilingual content'],
        ['Chinese diaspora', 'US/Canada/UK/Australia', 'Bilingual', 'Identity, cross-cultural health, dual-language'],
        ['Students', 'Global', 'Mandarin + English', 'Career advice, health habits, financial literacy'],
    ]
)


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 4 — SOCIAL MEDIA STRATEGY
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('4. Social Media Strategy', level=1)

add_para(
    'Social media serves three functions for Joyous: (1) driving listeners to the podcast, '
    '(2) building Jodi\'s personal brand as a cross-cultural authority, and (3) generating '
    'leads for the placement agency business. Each platform requires a distinct approach.'
)

# -- Platform Tactics
doc.add_heading('Platform-Specific Tactics', level=2)

doc.add_heading('LinkedIn', level=3)
add_bold_para('Role: Professional authority + placement agency lead generation', color=WARM_GOLD)

add_bullet('Post 4-5x per week: mix of original insights, episode clips, and industry commentary')
add_bullet('Write 1-2 long-form LinkedIn articles per month (these rank in Google search)')
add_bullet('Comment thoughtfully on 10-15 posts per day from LPs, GPs, and allocators')
add_bullet('Share episode clips natively (LinkedIn video gets 3x more reach than links)')
add_bullet('Build a "Venture Capital LP" and "Health Tech" focused newsletter on LinkedIn')
add_bullet('Connect with 20-30 new relevant professionals per week')
add_bullet('Content themes: fundraising insights, LP perspectives, cross-border capital flow, health-tech trends')

doc.add_heading('Twitter/X', level=3)
add_bold_para('Role: Real-time commentary + venture capital community engagement', color=WARM_GOLD)

add_bullet('Tweet 3-5x daily: threads, quote tweets, episode clips, hot takes')
add_bullet('Create weekly threads: "5 things I learned from [guest name] on Joyous this week"')
add_bullet('Engage with VC Twitter: reply to @HarryStebbings, @jason, @chaaboravsky, @balaboravsky')
add_bullet('Use Twitter Spaces for live bilingual Q&A sessions (monthly)')
add_bullet('Pin a tweet linking to the latest episode with a compelling hook')

doc.add_heading('Instagram', level=3)
add_bold_para('Role: Lifestyle brand + health content + behind-the-scenes', color=WARM_GOLD)

add_bullet('Post 4-5x per week: Reels (repurposed Shorts), carousels, Stories')
add_bullet('Reels: 30-60 sec clips from episodes, health tips, "day in the life" content')
add_bullet('Carousels: "3 wealth habits from [guest]", "East vs. West health approaches"')
add_bullet('Stories: Daily behind-the-scenes, polls, Q&A, podcast recording clips')
add_bullet('Use Link in Bio tool (Linktree or Stan Store) to drive to latest episode')
add_bullet('Aesthetic: Warm tones, gold accents, clean typography\u2014match the Joyous brand')

doc.add_heading('TikTok', level=3)
add_bold_para('Role: Discovery engine + viral growth + younger audience acquisition', color=WARM_GOLD)

add_bullet('Post 1-2x daily: repurpose YouTube Shorts + create TikTok-native content')
add_bullet('Hook viewers in the first 1.5 seconds (text overlay + surprising statement)')
add_bullet('Series ideas: "Rich vs. Wealthy", "Health Myths Debunked", "VC Explained in 60 Seconds"')
add_bullet('Use trending sounds when they align with content (don\'t force it)')
add_bullet('Engage with comments aggressively in the first hour after posting')
add_bullet('Cross-post top performers to Douyin (\u6296\u97f3) with Chinese captions')

# -- Content Calendar
doc.add_heading('Weekly Content Calendar Template', level=2)

add_styled_table(
    ['Day', 'LinkedIn', 'Twitter/X', 'Instagram', 'TikTok'],
    [
        ['Mon', 'Industry insight post', '3 tweets + thread', 'Reel (Shorts)', 'Short clip'],
        ['Tue', 'Episode announcement', 'Episode thread', 'Carousel (guest quotes)', 'Episode clip'],
        ['Wed', 'Engagement day (comments)', '5 quote tweets', 'Stories (BTS)', 'Trending format'],
        ['Thu', 'Article or newsletter', '3 tweets', 'Reel (health tip)', 'Health myth clip'],
        ['Fri', 'Weekend reading rec', 'Casual tweet', 'Carousel (weekly recap)', 'Fun/lifestyle'],
        ['Sat', '\u2014', 'Weekend engagement', 'Stories (personal)', '\u2014'],
        ['Sun', '\u2014', '\u2014', 'Stories (week preview)', '\u2014'],
    ]
)

# -- Brand Positioning
doc.add_heading('Personal Brand Positioning', level=2)

add_para(
    'Jodi operates across two worlds that must be carefully balanced: the entertainment/media '
    'brand (Joyous) and the professional/financial services brand (placement agency). These '
    'should complement, not conflict.'
)

doc.add_heading('Split Strategy: Two Brands, One Person', level=3)

add_styled_table(
    ['Dimension', 'Joyous Brand', 'Professional/Placement Brand'],
    [
        ['Tone', 'Warm, curious, accessible', 'Authoritative, precise, institutional'],
        ['Platforms', 'YouTube, Instagram, TikTok, Xiaohongshu', 'LinkedIn, Twitter/X, WeChat'],
        ['Content', 'Health tips, lifestyle, guest stories', 'Market analysis, fundraising, LP insights'],
        ['Visual', 'Gold/teal, warm photography', 'Navy/white, clean corporate'],
        ['CTA', '"Subscribe / Listen / Join"', '"Connect / Schedule a call / Learn more"'],
        ['Audience', 'Health-curious, culturally curious', 'GPs, LPs, allocators, fund managers'],
    ]
)

add_callout_box(
    'LinkedIn is the bridge platform. Jodi should post both Joyous episode content AND '
    'professional placement insights on LinkedIn. This creates natural crossover: an LP who '
    'discovers a Joyous episode about longevity investing may then learn about Jodi\'s '
    'placement services. The podcast is the top of the funnel; LinkedIn is the conversion layer.',
    title='KEY INSIGHT:'
)

# -- Engagement Tactics
doc.add_heading('Engagement Tactics', level=2)

add_numbered('Reply to every comment within 1 hour on all platforms for the first 6 months', bold_prefix='Comment-first: ')
add_numbered('DM every new follower with a personalized welcome (use ManyChat for automation)', bold_prefix='DM welcomes: ')
add_numbered('Create polls and "this or that" Stories to drive engagement metrics', bold_prefix='Interactive content: ')
add_numbered('Host monthly Instagram Lives or Twitter Spaces with a guest from a recent episode', bold_prefix='Live events: ')
add_numbered('Tag guests in all posts and ask them to reshare (provide pre-written copy)', bold_prefix='Guest amplification: ')
add_numbered('Create a "listener of the month" feature to reward engaged community members', bold_prefix='Community rewards: ')

# -- Hashtag Strategy
doc.add_heading('Hashtag Strategies', level=2)

doc.add_heading('English Hashtags', level=3)
add_para(
    '#JoyousPodcast #HealthAndWealth #WellnessInvesting #VentureCapital #Longevity '
    '#BilingualPodcast #EastMeetsWest #HealthTech #WealthMindset #AngelInvesting '
    '#WomenInVC #AsianAmerican #PodcastRecommendation #FinancialWellness #Biohacking'
)

doc.add_heading('Chinese Hashtags', level=3)
add_para(
    '#\u559c\u60a6\u64ad\u5ba2 #\u5065\u5eb7\u4e0e\u8d22\u5bcc #\u98ce\u9669\u6295\u8d44 #\u5065\u5eb7\u79d1\u6280 #\u534e\u4eba\u521b\u4e1a #\u4e2d\u7f8e\u5bf9\u6bd4 '
    '#\u6295\u8d44\u7b14\u8bb0 #\u7406\u8d22\u77e5\u8bc6 #\u5973\u6027\u529b\u91cf #\u5973\u6027\u6295\u8d44\u4eba #\u53cc\u8bed\u64ad\u5ba2 '
    '#\u7f8e\u56fd\u751f\u6d3b #\u5065\u5eb7\u751f\u6d3b #\u8d22\u5bcc\u81ea\u7531 #\u521b\u4e1a\u6545\u4e8b'
)


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 5 — PRIME LAB SHOW CONCEPT
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('5. Prime Lab Show Concept', level=1)

add_para(
    'Prime Lab is a high-concept entertainment format: company executives pitch their businesses '
    'while riding alongside a professional F1 driver at race speeds. It combines the adrenaline '
    'of motorsport with the intellectual stakes of venture pitching\u2014a format with no direct '
    'competitor in the market.'
)

# -- Show Format
doc.add_heading('Show Format Details', level=2)

doc.add_heading('Episode Structure', level=3)

add_numbered('Cold open with the most dramatic moment from the ride (15 seconds)', bold_prefix='Teaser: ')
add_numbered('Guest introduction, company overview, and why they\'re here (2 minutes)', bold_prefix='Meet the Pitcher: ')
add_numbered('Jodi interviews the exec about their business thesis before they get in the car (5 minutes)', bold_prefix='Pre-Ride Brief: ')
add_numbered('The exec pitches while riding shotgun with an F1 driver at speed (8-12 minutes)', bold_prefix='The Hot Lap Pitch: ')
add_numbered('Jodi and a panel discuss what they heard; guest reflects on the experience (5 minutes)', bold_prefix='Debrief: ')
add_numbered('Key business metrics, contact info, and a "would you invest?" audience poll (2 minutes)', bold_prefix='The Verdict: ')

add_callout_box(
    'Total episode length: 25-35 minutes. This is intentionally shorter than typical podcast '
    'interviews\u2014the F1 format demands intensity over length. The pitch-under-pressure format '
    'reveals authentic personality in a way that boardroom pitches never can.',
    title='FORMAT NOTE:'
)

doc.add_heading('The "Hot Lap Pitch" Rules', level=3)
add_bullet('Exec must pitch continuously while the car is moving\u2014no pauses allowed')
add_bullet('Driver maintains race pace; no slowing down for the pitcher')
add_bullet('Helmet-cam and in-car audio capture every moment')
add_bullet('Heart rate monitor displayed on screen (viewers see stress levels in real time)')
add_bullet('If the exec stops talking for more than 5 seconds, a buzzer sounds')

# -- Production Requirements
doc.add_heading('Production Requirements', level=2)

doc.add_heading('Venue & Track', level=3)
add_bullet('Partner with an F1 experience provider (e.g., F1 Experiences, Circuit of the Americas)')
add_bullet('Alternative: Professional racing circuits that allow media access (Laguna Seca, Silverstone)')
add_bullet('Budget option for pilot: Exotic car racing experiences (SpeedVegas, Exotics Racing)')
add_bullet('Minimum: closed circuit, professional driver, safety crew, insurance')

doc.add_heading('Equipment Needed', level=3)
add_styled_table(
    ['Item', 'Purpose', 'Est. Cost'],
    [
        ['GoPro Hero 12 (x4)', 'In-car cameras (dash, driver, passenger, rear)', '$2,000'],
        ['DJI Pocket 3', 'Stabilized walk-and-talk pre/post ride', '$500'],
        ['Wireless lav mics (x2)', 'In-car audio for driver and pitcher', '$600'],
        ['Heart rate monitors (x2)', 'Real-time biometric overlay', '$200'],
        ['Drone (DJI Mini 4 Pro)', 'Aerial track shots', '$800'],
        ['Track rental (half day)', 'Private session for filming', '$5,000-15,000'],
        ['F1 / pro driver fee', 'Professional driver for hot laps', '$3,000-10,000'],
        ['Insurance', 'Liability coverage for filming', '$2,000-5,000'],
        ['Post-production', 'Editing, graphics, sound design per episode', '$3,000-5,000'],
    ]
)

add_bold_para('Estimated pilot episode budget: $20,000-45,000', size=12, color=DEEP_TEAL)
add_bold_para('Estimated per-episode budget (after pilot): $15,000-30,000', size=12, color=DEEP_TEAL)

# -- Sponsorship
doc.add_heading('Sponsorship Package Ideas', level=2)

add_styled_table(
    ['Tier', 'Name', 'Price', 'Includes'],
    [
        ['Platinum', 'Title Sponsor', '$50,000/season', 'Logo on car, all episodes, presenting sponsor credit, 2 exec spots'],
        ['Gold', 'Track Sponsor', '$25,000/season', 'Logo on track signage, 3 episodes, 1 exec spot'],
        ['Silver', 'Lap Sponsor', '$10,000/episode', 'In-episode mention, logo overlay, social media feature'],
        ['Bronze', 'Pit Crew', '$5,000/episode', 'End-card logo, social media mention'],
    ]
)

doc.add_heading('Natural Sponsor Categories', level=3)
add_bullet('Luxury automotive brands (Porsche, McLaren, Ferrari\u2014product placement)')
add_bullet('Fintech and banking (Morgan Stanley, Goldman Sachs\u2014aspirational alignment)')
add_bullet('Health and wellness brands (Whoop, Oura Ring\u2014biometric tie-in)')
add_bullet('Enterprise SaaS (pitch tool integration, CRM sponsors)')
add_bullet('Luxury watch brands (motorsport heritage: TAG Heuer, Rolex, IWC)')

# -- Distribution Plan
doc.add_heading('Distribution Plan', level=2)

add_bullet('YouTube as primary platform (visual format demands video-first)', bold_prefix='YouTube: ')
add_bullet('Full episodes on Spotify Video, Apple Podcasts (audio version)', bold_prefix='Podcast: ')
add_bullet('30-second pitch highlight clips on TikTok, Instagram Reels, YouTube Shorts', bold_prefix='Short-form: ')
add_bullet('Bilibili for Chinese market (F1 is massive in China)', bold_prefix='China: ')
add_bullet('Pitch to streaming services (Netflix, Amazon Prime) after 6-8 episodes prove concept', bold_prefix='Streaming: ')

# -- Prime Movers Lab Integration
doc.add_heading('Integration with Prime Movers Lab', level=2)

add_para(
    'As a venture partner at Prime Movers Lab, Jodi has a natural pipeline of portfolio companies '
    'that could serve as pitchers on Prime Lab. This creates a virtuous cycle:'
)

add_bullet('PML portfolio companies get media exposure and brand building')
add_bullet('The show demonstrates PML\'s portfolio quality to potential LPs')
add_bullet('Jodi\'s involvement connects her placement agency to PML\'s fundraising')
add_bullet('PML\'s deep-tech focus (energy, biotech, transportation) provides compelling pitch material')
add_bullet('PML founder Dakin Sloss could serve as a recurring judge/commentator')

# -- Event Hosting
doc.add_heading('Event Hosting Strategy', level=2)

add_bullet('Host live "Prime Lab" events at major conferences (CES, SXSW, Web Summit)', bold_prefix='Conference tie-ins: ')
add_bullet('Create "Prime Lab Track Day" experiences for LPs and portfolio company execs', bold_prefix='LP events: ')
add_bullet('Partner with F1 races for viewing parties with live pitch competitions', bold_prefix='F1 race events: ')
add_bullet('Annual "Prime Lab Championship" with past pitchers competing for audience favorite', bold_prefix='Season finale: ')


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 6 — PLACEMENT AGENCY STRATEGY
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('6. Placement Agency Business Strategy', level=1)

add_para(
    'The placement agency represents the most significant revenue opportunity in Jodi\'s ecosystem. '
    'By connecting venture fund managers (GPs) with limited partners (LPs) across global markets, '
    'the agency leverages Jodi\'s unique cross-border network and the Joyous podcast\'s relationship-'
    'building power.'
)

# -- Business Structure
doc.add_heading('Business Structure Recommendations', level=2)

add_styled_table(
    ['Structure', 'Pros', 'Cons', 'Recommendation'],
    [
        ['LLC (Single Member)', 'Simple, low cost, pass-through tax', 'Less institutional credibility', 'Start here'],
        ['LLC (Multi-Member)', 'Allows partners, flexible structure', 'Operating agreement needed', 'If adding partners'],
        ['S-Corp Election', 'Payroll tax savings above $80K income', 'More admin, payroll required', 'Once profitable'],
        ['C-Corp', 'Institutional LPs prefer; VC-backable', 'Double taxation', 'Only if raising capital'],
    ]
)

add_callout_box(
    'Start as a Delaware LLC with S-Corp tax election once placement fee revenue exceeds $80,000/year. '
    'Delaware provides the most favorable legal framework for financial services firms and is expected '
    'by institutional LPs. Register as a foreign LLC in your state of residence.',
    title='RECOMMENDATION:'
)

# -- Licensing Requirements
doc.add_heading('Licensing & Regulatory Requirements', level=2)

doc.add_heading('United States', level=3)

add_styled_table(
    ['Requirement', 'Details', 'Cost', 'Timeline'],
    [
        ['SEC Broker-Dealer Registration', 'Required if receiving transaction-based compensation for securities placement', '$5,000-15,000 (legal)', '3-6 months'],
        ['FINRA Series 82', 'Private securities offerings representative exam', '$305 exam fee', '2-3 months study'],
        ['FINRA Series 63', 'State securities law (if required by state)', '$147 exam fee', '1-2 months study'],
        ['State Registration', 'Register in states where you solicit LPs', '$200-500/state', '2-4 weeks'],
        ['SEC Form BD', 'Broker-dealer registration form', 'Included in legal', '3-6 months'],
    ]
)

add_callout_box(
    'CRITICAL: Receiving transaction-based compensation (% of capital raised) for introducing LPs to '
    'fund managers almost certainly requires broker-dealer registration with the SEC and FINRA membership. '
    'Operating without registration carries severe penalties. Consult a securities attorney before '
    'accepting your first placement fee. Some exemptions may apply for M&A advisory (see SEC no-action '
    'letters), but fund placement is generally NOT exempt.',
    title='LEGAL WARNING:'
)

doc.add_heading('International Considerations', level=3)

add_styled_table(
    ['Jurisdiction', 'Regulator', 'Key Requirement', 'Notes'],
    [
        ['United Kingdom', 'FCA', 'FCA Authorization or Appointed Representative', 'Required for UK LP solicitation'],
        ['Singapore', 'MAS', 'Capital Markets Services License', 'Required for SG-based LP solicitation'],
        ['Hong Kong', 'SFC', 'Type 1 License', 'Required for HK LP solicitation'],
        ['Middle East (UAE)', 'DFSA/ADGM', 'Varies by emirate', 'DIFC or ADGM registration'],
        ['Australia', 'ASIC', 'AFS License', 'Required for AU LP solicitation'],
    ]
)

add_para(
    'For initial operations, consider partnering with a licensed placement agent in each jurisdiction '
    'as an "introducing broker" or under their license umbrella, rather than obtaining your own '
    'licenses in every market.'
)

# -- Target LP Regions
doc.add_heading('Target LP Regions & Approach Strategies', level=2)

doc.add_heading('Tier 1: Immediate Targets (Months 1-6)', level=3)

add_bullet('Strong family office culture; less regulated; relationship-driven', bold_prefix='Middle East (UAE, Saudi, Qatar): ')
add_bullet('Familiar territory; large pension funds, family offices, sovereign wealth', bold_prefix='Asia (Singapore, Hong Kong, Taipei): ')
add_bullet('Superannuation funds seeking venture exposure; under-allocated to VC', bold_prefix='Australia: ')

doc.add_heading('Tier 2: Medium-Term (Months 6-12)', level=3)

add_bullet('Energy wealth seeking diversification; UTIMCO, Texas TRS, family offices', bold_prefix='Texas: ')
add_bullet('Auto industry sovereign wealth funds; Porsche Ventures, BMW iVentures connections', bold_prefix='Germany (Stuttgart/Munich): ')
add_bullet('Government pension funds, sovereign wealth (Norges Bank, AP funds)', bold_prefix='Nordics: ')
add_bullet('Large pension funds (CPPIB, CDPQ, OTPP); sophisticated VC allocators', bold_prefix='Canada: ')

doc.add_heading('Tier 3: Long-Term (Year 1+)', level=3)

add_bullet('GPIF, corporate pension funds, trading houses (Softbank, Mitsui)', bold_prefix='Japan: ')
add_bullet('National Pension Service, Samsung Ventures, Hyundai connections', bold_prefix='Korea: ')
add_bullet('Massive capital pools; regulatory complexity; requires local partners', bold_prefix='Mainland China: ')

# -- Revenue Model
doc.add_heading('Revenue Model', level=2)

add_styled_table(
    ['Revenue Stream', 'Rate', 'Example', 'Annual Potential'],
    [
        ['Placement Fee', '1.0-2.0% of capital raised', '$100M fund = $1-2M fee', '$500K-5M'],
        ['Retainer Fee', '$5,000-15,000/month', '3-5 GP clients on retainer', '$180K-900K'],
        ['Success Bonus', '0.25-0.50% on oversubscription', 'If fund exceeds target', 'Variable'],
        ['Ongoing Advisory', '$2,000-5,000/month', 'Post-close LP relationship management', '$72K-180K'],
        ['Event Introduction Fee', '$1,000-5,000/intro', 'Conference LP introductions', '$50K-200K'],
    ]
)

add_callout_box(
    'Industry standard placement fees range from 1.0% to 2.5% of capital committed. For a new '
    'placement agent, start at 1.5% with a minimum fee floor of $100,000 per mandate. As track '
    'record builds, increase to 2.0%. Retainer fees of $5,000-10,000/month against placement '
    'fees provide cash flow during the typically 6-18 month fundraising cycle.',
    title='PRICING GUIDANCE:'
)

# -- Competitive Reference
doc.add_heading('Competitive Landscape: Mountside Ventures as Reference', level=2)

add_para(
    'Mountside Ventures (London-based) is the most relevant competitive benchmark. Founded by '
    'former GPs, they provide placement and fundraising advisory for emerging managers. Key '
    'takeaways from their model:'
)

add_bullet('They focus on emerging and first-time managers ($20M-200M fund sizes)\u2014an underserved niche')
add_bullet('They combine placement with LP introductions, pitch deck review, and fundraising strategy')
add_bullet('Their content marketing (blog, LinkedIn, podcast appearances) drives 40%+ of inbound leads')
add_bullet('They maintain a proprietary LP database of 3,000+ institutional investors')
add_bullet('Team of 5-8 people; lean operations with high margins')

add_callout_box(
    'Jodi\'s advantages over Mountside: (1) bilingual access to Chinese/Asian LPs that Mountside '
    'cannot reach, (2) the Joyous podcast as a content engine and trust builder, (3) Prime Movers '
    'Lab affiliation providing deep-tech credibility, and (4) personal network across 12+ cities '
    'and 3 continents.',
    title='COMPETITIVE EDGE:'
)

# -- Podcast-to-Pipeline
doc.add_heading('How the Podcast Feeds the Placement Agency Pipeline', level=2)

add_para('The Joyous podcast is not just a media product\u2014it is the placement agency\'s primary '
         'business development tool. Here is the conversion funnel:')

add_numbered('Podcast guests include fund managers (GPs) who need placement services', bold_prefix='Top of Funnel: ')
add_numbered('Episode discussion reveals fund thesis, team quality, and market opportunity', bold_prefix='Qualification: ')
add_numbered('Post-episode, Jodi follows up with "I know LPs who would love this thesis"', bold_prefix='Warm Introduction: ')
add_numbered('GP engages placement agency for formal fundraising mandate', bold_prefix='Engagement: ')
add_numbered('LP listeners discover funds through the podcast; contact Jodi for access', bold_prefix='Inbound LP Interest: ')

add_callout_box(
    'Target: 20% of podcast guests should be potential GP clients for the placement agency. '
    'This means 1 in 5 episodes should feature an emerging fund manager. These episodes also '
    'serve as valuable content for LP listeners evaluating the venture landscape.',
    title='TARGET RATIO:'
)


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 7 — MONETIZATION & REVENUE STREAMS
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('7. Monetization & Revenue Streams', level=1)

add_para(
    'Joyous sits at the center of a multi-revenue ecosystem. The podcast itself generates modest '
    'direct revenue, but its real value is as a force multiplier for higher-margin businesses: '
    'placement fees, consulting, events, and brand partnerships.'
)

# -- Revenue Map
doc.add_heading('Revenue Stream Overview', level=2)

add_styled_table(
    ['Stream', 'Revenue Potential (Year 1)', 'Margin', 'Difficulty', 'Timeline'],
    [
        ['Placement Agency Fees', '$200K-2M', '70-85%', 'High (licensing)', 'Months 3-12'],
        ['Podcast Sponsorships', '$20K-80K', '90%', 'Medium', 'Months 3-6'],
        ['Speaking Engagements', '$30K-100K', '95%', 'Medium', 'Months 1-6'],
        ['Event Hosting/Tickets', '$15K-50K', '40-60%', 'Medium-High', 'Months 6-12'],
        ['Consulting/Advisory', '$50K-200K', '90%', 'Low', 'Months 1-3'],
        ['YouTube Revenue', '$5K-30K', '70%', 'Low', 'Months 4-12'],
        ['Premium Membership', '$10K-40K', '85%', 'Medium', 'Months 4-8'],
        ['Brand Partnerships', '$20K-100K', '80%', 'Medium', 'Months 6-12'],
    ]
)

add_bold_para('Total Year 1 Revenue Potential: $350,000 - $2,600,000', size=14, color=DEEP_TEAL)

# -- Sponsorships
doc.add_heading('Podcast Sponsorships', level=2)

add_para(
    'Podcast sponsorship rates are typically measured in CPM (cost per thousand downloads). '
    'For niche, high-value audiences like Joyous\'s (investors, executives, health-conscious '
    'professionals), CPMs run significantly higher than average.'
)

add_styled_table(
    ['Placement', 'Industry Avg CPM', 'Joyous Target CPM', 'Notes'],
    [
        ['Pre-roll (15 sec)', '$18-25', '$30-40', 'Before episode starts'],
        ['Mid-roll (60 sec)', '$25-40', '$50-75', 'Highest value; host-read'],
        ['Post-roll (30 sec)', '$10-15', '$15-25', 'End of episode'],
        ['Custom Integration', 'N/A', '$1,000-3,000/ep', 'Guest interview with sponsor exec'],
    ]
)

doc.add_heading('Ideal Sponsor Categories', level=3)
add_bullet('Wealth management and fintech (Wealthfront, Betterment, Personal Capital)')
add_bullet('Health and wellness (Athletic Greens/AG1, Whoop, Eight Sleep, Momentous)')
add_bullet('Luxury travel (Cathay Pacific, Singapore Airlines, Mandarin Oriental)')
add_bullet('Executive education (Masterclass, Coursera, Wharton Executive Ed)')
add_bullet('B2B SaaS for investors (Carta, AngelList, Affinity CRM)')

# -- Premium Content
doc.add_heading('Premium Content / Membership', level=2)

add_para(
    'Launch a membership program after establishing a consistent free audience (target: after '
    '20 episodes and 5,000+ combined subscribers across platforms).'
)

doc.add_heading('Membership Tiers', level=3)

add_styled_table(
    ['Tier', 'Price', 'Includes'],
    [
        ['Free', '$0', 'All regular episodes, YouTube content, social media'],
        ['Joyous Insider', '$9.99/month', 'Extended interviews, bonus episodes, early access, private Discord'],
        ['Joyous Circle', '$49.99/month', 'Above + monthly live Q&A with Jodi, curated health/wealth newsletter'],
        ['Joyous Patron', '$199/month', 'Above + quarterly 1-on-1 call, invitation to live events, LP/GP introductions'],
    ]
)

# -- Events
doc.add_heading('Event Hosting Revenue', level=2)

add_bullet('Quarterly "Joyous Salon" dinners: 20-30 curated guests, $150-500/ticket', bold_prefix='Intimate dinners: ')
add_bullet('Annual "Joyous Summit": 200-500 attendees, $500-1,500/ticket, sponsor-supported', bold_prefix='Annual conference: ')
add_bullet('Prime Lab live filming events: 50-100 VIP attendees, $250-750/ticket', bold_prefix='Prime Lab events: ')
add_bullet('Bilingual networking events in major cities (NYC, SF, London, Singapore, Shanghai)', bold_prefix='Cross-cultural mixers: ')

# -- Consulting
doc.add_heading('Consulting & Advisory Revenue', level=2)

add_bullet('GP fundraising strategy consulting: $5,000-15,000 per engagement', bold_prefix='Fund strategy: ')
add_bullet('LP introduction facilitation: $2,000-5,000 per warm introduction', bold_prefix='LP introductions: ')
add_bullet('Cross-border investment advisory: $500/hour for East-West bridge consulting', bold_prefix='Cross-border advisory: ')
add_bullet('Podcast launch consulting: $3,000-8,000 for executives launching their own shows', bold_prefix='Media consulting: ')

# -- Speaking
doc.add_heading('Speaking Engagements', level=2)

add_para('Jodi\'s unique background commands premium speaking fees at industry events:')

add_styled_table(
    ['Event Type', 'Fee Range', 'Examples'],
    [
        ['Industry conference keynote', '$5,000-15,000', 'SuperReturn, ILPA Summit, AVCJ'],
        ['Corporate event / retreat', '$3,000-10,000', 'Family office retreats, bank events'],
        ['University/MBA lecture', '$1,000-3,000', 'Wharton, Harvard, INSEAD'],
        ['Panel moderation', '$2,000-5,000', 'LP/GP conferences, health-tech events'],
        ['Private dinner speaker', '$3,000-8,000', 'Exclusive investor gatherings'],
    ]
)

# -- Brand Partnerships
doc.add_heading('Brand Partnerships', level=2)

add_bullet('Long-term ambassador deals with luxury/wellness brands (6-12 month contracts)', bold_prefix='Ambassador deals: ')
add_bullet('Co-branded content series with health tech companies', bold_prefix='Content partnerships: ')
add_bullet('Sponsored travel/experience content for luxury hospitality brands', bold_prefix='Travel partnerships: ')
add_bullet('Chinese brand partnerships for cross-border content (unique value proposition)', bold_prefix='China bridge deals: ')


# ═══════════════════════════════════════════════════════════════════════
#  CHAPTER 8 — 90-DAY LAUNCH PLAN
# ═══════════════════════════════════════════════════════════════════════
add_page_break()

doc.add_heading('8. 90-Day Launch Plan', level=1)

add_para(
    'This chapter provides a week-by-week action plan for the first 90 days of the Joyous '
    'growth initiative. The plan assumes 6 edited episodes are ready for release and focuses '
    'on distribution, audience building, and business development.'
)

# -- Weeks 1-2
doc.add_heading('Phase 1: Foundation (Weeks 1-2)', level=2)

add_bold_para('Theme: "Set the stage"', color=WARM_GOLD)

doc.add_heading('Week 1', level=3)
add_bullet('Set up podcast hosting (Buzzsprout or Transistor) and submit RSS to all platforms')
add_bullet('Create YouTube channel with bilingual branding, banner, and trailer')
add_bullet('Set up Xiaohongshu account and WeChat Official Account')
add_bullet('Optimize LinkedIn profile for dual positioning (Joyous + placement)')
add_bullet('Upload first 3 episodes across all platforms (staggered: Mon/Wed/Fri)')
add_bullet('Create 5 YouTube Shorts from existing episode content')
add_bullet('Design thumbnail templates (3 variations for A/B testing)')
add_bullet('Set up link-in-bio page (Linktree or Stan Store)')

doc.add_heading('Week 2', level=3)
add_bullet('Upload episodes 4-6 (complete the initial catalog)')
add_bullet('Submit to Apple Podcasts, Spotify, Amazon Music, iHeartRadio, Pocket Casts')
add_bullet('Register on Ximalaya and upload Chinese recap of episodes 1-3')
add_bullet('Begin daily LinkedIn posting (1 post/day + 10 comments on others\' posts)')
add_bullet('Launch Twitter/X with 5 tweets/day cadence')
add_bullet('Create Instagram account; post 3 Reels from episode clips')
add_bullet('Send personal emails to 50 contacts asking for Apple Podcasts reviews')
add_bullet('Begin guest outreach for next 6 episodes (target: 10 pitches sent)')

# -- Weeks 3-4
doc.add_heading('Phase 2: Acceleration (Weeks 3-4)', level=2)

add_bold_para('Theme: "Build momentum"', color=WARM_GOLD)

doc.add_heading('Week 3', level=3)
add_bullet('Launch TikTok account; post 2 clips/day from existing content')
add_bullet('Record first new episode post-launch (guest secured in Week 2)')
add_bullet('Create first full Chinese episode or Chinese recap video for Bilibili')
add_bullet('Launch Xiaohongshu content calendar: 4 posts this week')
add_bullet('Write first WeChat Official Account article (bilingual episode summary)')
add_bullet('Reach out to 5 complementary podcasters for guest swap opportunities')
add_bullet('Begin placement agency legal consultation (securities attorney)')

doc.add_heading('Week 4', level=3)
add_bullet('Analyze first 3 weeks of analytics: identify top-performing content formats')
add_bullet('Double down on highest-performing platform and content type')
add_bullet('Record episode #2 post-launch')
add_bullet('Create a media kit for sponsorship outreach (1-pager + rate card)')
add_bullet('Begin sponsor outreach to 10 potential sponsors')
add_bullet('Establish weekly content production workflow (batch record, batch edit)')
add_bullet('Apply for YouTube Partner Program (if eligible)')

# -- Weeks 5-8
doc.add_heading('Phase 3: Growth (Weeks 5-8)', level=2)

add_bold_para('Theme: "Scale what works"', color=WARM_GOLD)

add_bullet('Maintain weekly episode release schedule (1 new episode per week)')
add_bullet('Maintain 4-5 YouTube Shorts per week')
add_bullet('First guest collaboration episode (appear on another podcast)')
add_bullet('Launch Joyous Discord or community channel for engaged listeners')
add_bullet('First sponsored episode or brand integration')
add_bullet('Begin FINRA Series 82 exam preparation')
add_bullet('Attend 1-2 industry events and promote Joyous/collect LP contacts')
add_bullet('First live Twitter Space or Instagram Live event')
add_bullet('Create "best of" compilation video for YouTube')
add_bullet('Hire bilingual content assistant (part-time) for Chinese platform management')
add_bullet('Begin Prime Lab pilot episode pre-production (venue scouting, driver booking)')
add_bullet('Reach 100 Apple Podcasts reviews')

# -- Weeks 9-12
doc.add_heading('Phase 4: Monetization (Weeks 9-12)', level=2)

add_bold_para('Theme: "Turn audience into revenue"', color=WARM_GOLD)

add_bullet('Close first podcast sponsor deal ($500-2,000 per episode)')
add_bullet('Launch membership/Patreon with Insider and Circle tiers')
add_bullet('Host first "Joyous Salon" dinner event (20-30 attendees)')
add_bullet('Onboard first 1-2 GP clients for placement agency (if licensing permits)')
add_bullet('Secure first paid speaking engagement')
add_bullet('Film Prime Lab pilot episode')
add_bullet('Reach 1,000+ YouTube subscribers')
add_bullet('Reach 2,500+ combined podcast subscribers')
add_bullet('Establish regular bilingual content pipeline (EN + CN weekly)')
add_bullet('Create Q2 strategy based on Q1 data and learnings')

# -- Priority Matrix
doc.add_heading('Priority Matrix', level=2)

add_styled_table(
    ['Priority', 'Action', 'Impact', 'Effort'],
    [
        ['P0 (Do First)', 'Distribute to all podcast platforms', 'Critical', 'Low'],
        ['P0', 'Launch YouTube with Shorts-first strategy', 'Critical', 'Medium'],
        ['P0', 'Daily LinkedIn posting', 'High', 'Low'],
        ['P0', 'Consult securities attorney re: placement licensing', 'Critical', 'Low'],
        ['P1', 'Launch Chinese platforms (Ximalaya, Xiaohongshu)', 'High', 'Medium'],
        ['P1', 'Begin guest outreach pipeline', 'High', 'Medium'],
        ['P1', 'Create media kit for sponsors', 'High', 'Low'],
        ['P2', 'Launch TikTok and Instagram', 'Medium', 'Medium'],
        ['P2', 'Begin FINRA exam prep', 'High', 'High'],
        ['P2', 'Prime Lab pre-production', 'Medium', 'High'],
        ['P3', 'Launch membership program', 'Medium', 'Medium'],
        ['P3', 'Host first live event', 'Medium', 'High'],
    ]
)

# -- Resource Needs
doc.add_heading('Resource Needs', level=2)

doc.add_heading('Team (Recommended Hires)', level=3)

add_styled_table(
    ['Role', 'Type', 'Monthly Cost', 'Priority'],
    [
        ['Podcast Editor', 'Freelance', '$800-1,500', 'Immediate'],
        ['Bilingual Content Assistant', 'Part-time', '$1,500-2,500', 'Month 2'],
        ['Social Media Manager', 'Part-time', '$1,500-3,000', 'Month 3'],
        ['Video Editor (Shorts/Clips)', 'Freelance', '$600-1,200', 'Month 1'],
        ['Thumbnail Designer', 'Freelance', '$200-400', 'Month 1'],
        ['Securities Attorney', 'Hourly ($400-600/hr)', '$2,000-5,000 (one-time)', 'Immediate'],
        ['Bookkeeper', 'Monthly', '$300-500', 'Month 3'],
    ]
)

# -- Budget
doc.add_heading('Budget Estimates (90-Day)', level=2)

add_styled_table(
    ['Category', 'Monthly', '90-Day Total', 'Notes'],
    [
        ['Podcast hosting', '$20', '$60', 'Buzzsprout or Transistor'],
        ['Editing (audio + video)', '$1,200', '$3,600', 'Freelance editor'],
        ['Video clips / Shorts', '$800', '$2,400', 'Freelance video editor'],
        ['Graphic design (thumbnails)', '$300', '$900', 'Freelance or Canva Pro'],
        ['Social media tools', '$100', '$300', 'Buffer/Hootsuite + Canva Pro'],
        ['Chinese platform management', '$500', '$1,500', 'Bilingual assistant (starting Month 2)'],
        ['Recording equipment (one-time)', '\u2014', '$2,000', 'Mic, camera, lighting, acoustic treatment'],
        ['Website (joyouspodcast.com)', '$30', '$90', 'Squarespace or WordPress'],
        ['Legal (securities consultation)', '\u2014', '$3,000-5,000', 'Initial consultation + filing prep'],
        ['Marketing / Ads', '$500', '$1,500', 'Instagram/YouTube ads for launch'],
        ['Travel (events)', '$1,000', '$3,000', '1-2 industry events'],
        ['Miscellaneous', '$300', '$900', 'Software, subscriptions, contingency'],
    ]
)

add_bold_para('Total 90-Day Budget Estimate: $19,250 - $21,250', size=14, color=DEEP_TEAL)
add_para(
    '(Excluding Prime Lab pilot episode production, which is budgeted separately at $20,000-45,000.)'
)

# ── Final page ─────────────────────────────────────────────────────────
add_page_break()

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('End of Strategy Document')
run.font.size = Pt(18)
run.font.color.rgb = DEEP_TEAL
run.font.name = 'Calibri'
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for Jodi Yang  |  March 2026')
run.font.size = Pt(12)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'This document is confidential and intended solely for the use of the named recipient. '
    'Do not distribute without permission.'
)
run.font.size = Pt(9)
run.font.color.rgb = MED_GRAY
run.font.name = 'Calibri'
run.font.italic = True

# ── Save ───────────────────────────────────────────────────────────────
output_path = '/Users/gregspero/joyous-podcast/Joyous_Podcast_Strategy.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
